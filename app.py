
# app_docx_borders_85pt_editable_v6_8_7_locked.py
import datetime, json, urllib.parse, urllib.request
from io import BytesIO

import matplotlib.pyplot as plt
import pandas as pd
import pytz
import streamlit as st
import swisseph as swe
from timezonefinder import TimezoneFinder

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

APP_TITLE = "DevoAstroBhav Kundali — Locked (v6.8.7)"
st.set_page_config(page_title=APP_TITLE, layout="wide", page_icon="🪔")

AYANAMSHA_VAL = swe.SIDM_LAHIRI
YEAR_DAYS     = 365.2422

# Fonts / sizing (slightly smaller base to avoid clipping)
BASE_FONT_PT = 8.0
LATIN_FONT = "Georgia"
HINDI_FONT = "Mangal"

HN = {'Su':'सूर्य','Mo':'चंद्र','Ma':'मंगल','Me':'बुध','Ju':'गुरु','Ve':'शुक्र','Sa':'शनि','Ra':'राहु','Ke':'केतु'}

def _apply_hindi_caption_style(paragraph, size_pt=11, underline=True, bold=True):
    if not paragraph.runs:
        paragraph.add_run("")
    r = paragraph.runs[0]
    r.bold = bold; r.underline = underline; r.font.size = Pt(size_pt)
    rpr = r._element.rPr or OxmlElement('w:rPr')
    if r._element.rPr is None: r._element.append(rpr)
    rfonts = rpr.find(qn('w:rFonts')) or OxmlElement('w:rFonts')
    if rpr.find(qn('w:rFonts')) is None: rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), HINDI_FONT)

def set_sidereal_locked():
    swe.set_sid_mode(AYANAMSHA_VAL, 0, 0)

def dms_exact(deg):
    d = int(deg); m_float = (deg - d) * 60.0; m = int(m_float); s = (m_float - m) * 60.0
    return d, m, s

def fmt_deg_sign(lon_sid):
    sign=int(lon_sid//30) + 1; deg_in_sign = lon_sid % 30.0
    d,m,s=dms_exact(deg_in_sign); s_rounded = int(round(s))
    if s_rounded == 60: s_rounded = 0; m += 1
    if m == 60: m = 0; d += 1; 
    if d == 30: d = 0
    return sign, f"{d:02d}°{m:02d}'{s_rounded:02d}\""

def kp_sublord(lon_sid):
    NAK=360.0/27.0
    ORDER = ['Ke','Ve','Su','Mo','Ma','Ra','Ju','Sa','Me']
    YEARS = {'Ke':7,'Ve':20,'Su':6,'Mo':10,'Ma':7,'Ra':18,'Ju':16,'Sa':19,'Me':17}
    part = lon_sid % 360.0; ni = int(part // NAK); pos = part - ni*NAK
    lord = ORDER[ni % 9]; start = ORDER.index(lord)
    seq = [ORDER[(start+i)%9] for i in range(9)]
    acc = 0.0
    for L in seq:
        seg = NAK * (YEARS[L]/120.0)
        if pos <= acc + seg + 1e-9: return lord, L
        acc += seg
    return lord, seq[-1]

def geocode(place, api_key):
    if not api_key: raise RuntimeError("Geoapify key missing. Add GEOAPIFY_API_KEY in Secrets.")
    base="https://api.geoapify.com/v1/geocode/search?"
    q = urllib.parse.urlencode({"text":place, "format":"json", "limit":1, "apiKey":api_key})
    with urllib.request.urlopen(base+q, timeout=15) as r: j = json.loads(r.read().decode())
    if j.get("results"):
        res=j["results"][0]; return float(res["lat"]), float(res["lon"]), res.get("formatted", place)
    raise RuntimeError("Place not found.")

def tz_from_latlon(lat, lon, dt_local):
    tf = TimezoneFinder(); tzname = tf.timezone_at(lat=lat, lng=lon) or "Etc/UTC"
    tz = pytz.timezone(tzname); dt_local_aware = tz.localize(dt_local)
    dt_utc_naive = dt_local_aware.astimezone(pytz.utc).replace(tzinfo=None)
    offset_hours = tz.utcoffset(dt_local_aware).total_seconds()/3600.0
    return tzname, offset_hours, dt_utc_naive

def sidereal_positions(dt_utc):
    jd = swe.julday(dt_utc.year, dt_utc.month, dt_utc.day, dt_utc.hour + dt_utc.minute/60 + dt_utc.second/3600)
    set_sidereal_locked(); flags = swe.FLG_SWIEPH | swe.FLG_SPEED | swe.FLG_SIDEREAL
    out = {}
    for code, p in [('Su',swe.SUN),('Mo',swe.MOON),('Ma',swe.MARS),('Me',swe.MERCURY),('Ju',swe.JUPITER),('Ve',swe.VENUS),('Sa',swe.SATURN)]:
        xx,_ = swe.calc_ut(jd, p, flags); out[code] = xx[0] % 360.0
    xx,_ = swe.calc_ut(jd, swe.MEAN_NODE, flags)  # Mean node locked
    out['Ra'] = xx[0] % 360.0; out['Ke'] = (out['Ra'] + 180.0) % 360.0
    ay = swe.get_ayanamsa_ut(jd); return jd, ay, out

def ascendant_sign(jd, lat, lon, ay):
    cusps, ascmc = swe.houses_ex(jd, lat, lon, b'P'); asc_trop = ascmc[0]; asc_sid = (asc_trop - ay) % 360.0
    return int(asc_sid // 30) + 1, asc_sid

def navamsa_sign_from_lon_sid(lon_sid):
    sign = int(lon_sid // 30) + 1; deg_in_sign = lon_sid % 30.0; pada = int(deg_in_sign // (30.0/9.0))
    if sign in (1,4,7,10): start = sign
    elif sign in (2,5,8,11): start = ((sign + 8 - 1) % 12) + 1
    else: start = ((sign + 4 - 1) % 12) + 1
    return ((start - 1 + pada) % 12) + 1

def positions_table_no_symbol(sidelons):
    rows=[]
    for code in ['Su','Mo','Ma','Me','Ju','Ve','Sa','Ra','Ke']:
        lon=sidelons[code]; sign, deg_str = fmt_deg_sign(lon); nak_lord, sub_lord = kp_sublord(lon)
        rows.append([HN[code], sign, deg_str, HN[nak_lord], HN[sub_lord]])
    return pd.DataFrame(rows, columns=["Planet","Sign","Degree","Nakshatra","Sub‑Nakshatra"])

ORDER = ['Ke','Ve','Su','Mo','Ma','Ra','Ju','Sa','Me']
YEARS = {'Ke':7,'Ve':20,'Su':6,'Mo':10,'Ma':7,'Ra':18,'Ju':16,'Sa':19,'Me':17}

def moon_balance_days(moon_sid):
    NAK=360.0/27.0; part = moon_sid % 360.0; ni = int(part // NAK); pos = part - ni*NAK
    md_lord = ORDER[ni % 9]; frac = pos/NAK; remaining_days = YEARS[md_lord]*(1 - frac)*YEAR_DAYS
    return md_lord, remaining_days

def build_mahadashas_days_utc(birth_utc_dt, moon_sid):
    md_lord, rem_days = moon_balance_days(moon_sid); end_limit = birth_utc_dt + datetime.timedelta(days=100*YEAR_DAYS)
    segments=[]; birth_md_start = birth_utc_dt; birth_md_end = min(birth_md_start + datetime.timedelta(days=rem_days), end_limit)
    segments.append({"planet": md_lord, "start": birth_md_start, "end": birth_md_end, "days": rem_days})
    idx = (ORDER.index(md_lord) + 1) % 9; t = birth_md_end
    while t < end_limit:
        L = ORDER[idx]; dur_days = YEARS[L]*YEAR_DAYS; end = min(t + datetime.timedelta(days=dur_days), end_limit)
        segments.append({"planet": L, "start": t, "end": end, "days": dur_days}); t = end; idx = (idx + 1) % 9
    return segments

def antar_segments_in_md_utc(md_lord, md_start_utc, md_days):
    res=[]; t=md_start_utc; start_idx=ORDER.index(md_lord)
    for i in range(9):
        L=ORDER[(start_idx+i)%9]; dur = YEARS[L]*(md_days/(120.0)); start = t; end = t + datetime.timedelta(days=dur)
        res.append((L, start, end, dur)); t = end
    return res

def pratyantars_in_antar_utc(antar_lord, antar_start_utc, antar_days):
    res=[]; t=antar_start_utc; start_idx=ORDER.index(antar_lord)
    for i in range(9):
        L=ORDER[(start_idx+i)%9]; dur = YEARS[L]*(antar_days/(120.0)); start = t; end = t + datetime.timedelta(days=dur)
        res.append((L, start, end)); t = end
    return res

def next_ant_praty_in_days_utc(now_utc, md_segments, days_window):
    rows=[]; horizon=now_utc + datetime.timedelta(days=days_window)
    for seg in md_segments:
        MD = seg["planet"]; ms = seg["start"]; me = seg["end"]; md_days = seg["days"]
        for AL, as_, ae, adays in antar_segments_in_md_utc(MD, ms, md_days):
            if ae < now_utc or as_ > horizon: continue
            for PL, ps, pe in pratyantars_in_antar_utc(AL, as_, adays):
                if pe < now_utc or ps > horizon: continue
                rows.append({"major":MD,"antar":AL,"pratyantar":PL,"end":pe})
    rows.sort(key=lambda r:r["end"]); return rows

def render_north_diamond(size_px=900, stroke=3):
    fig = plt.figure(figsize=(size_px/100, size_px/100), dpi=100)
    ax = fig.add_axes([0,0,1,1]); ax.axis('off')
    ax.plot([0.02,0.98,0.98,0.02,0.02],[0.02,0.02,0.98,0.98,0.02], linewidth=3, color='black')
    L,R,B,T = 0.02,0.98,0.02,0.98; ax.plot([L,R],[T,B],3,color='black'); ax.plot([L,R],[B,T],3,color='black')
    midL=(L,0.5); midR=(R,0.5); midT=(0.5,T); midB=(0.5,B)
    ax.plot([midL[0], midT[0]],[midL[1], midT[1]], linewidth=3, color='black')
    ax.plot([midT[0], midR[0]],[midT[1], midR[1]], linewidth=3, color='black')
    ax.plot([midR[0], midB[0]],[midR[1], midB[1]], linewidth=3, color='black')
    ax.plot([midB[0], midL[0]],[midB[1], midL[1]], linewidth=3, color='black')
    buf = BytesIO(); fig.savefig(buf, format='png', bbox_inches='tight', pad_inches=0.02); plt.close(fig); buf.seek(0); return buf

def rotated_house_labels(lagna_sign):
    order = [str(((lagna_sign - 1 + i) % 12) + 1) for i in range(12)]
    return {"1":order[0],"2":order[1],"3":order[2],"4":order[3],"5":order[4],"6":order[5],"7":order[6],"8":order[7],"9":order[8],"10":order[9],"11":order[10],"12":order[11]}

def kundali_w_p_with_centroid_labels(size_pt=220, lagna_sign=1):
    S=size_pt; TM=(S/2,0); RM=(S,S/2); BM=(S/2,S); LM=(0,S/2); P_lt=(S/4,S/4); P_rt=(3*S/4,S/4); P_rb=(3*S/4,3*S/4); P_lb=(S/4,3*S/4); O=(S/2,S/2)
    labels = rotated_house_labels(lagna_sign)
    houses = {"1":[TM,P_rt,O,P_lt],"2":[(0,0),TM,P_lt],"3":[(0,0),LM,P_lt],"4":[LM,O,P_lt,P_lb],"5":[LM,(0,S),P_lb],"6":[(0,S),BM,P_lb],"7":[BM,P_rb,O,P_lb],"8":[BM,(S,S),P_rb],"9":[RM,(S,S),P_rb],"10":[RM,O,P_rt,P_rb],"11":[(S,0),RM,P_rt],"12":[TM,(S,0),P_rt]}
    def centroid(poly):
        A=Cx=Cy=0.0; n=len(poly)
        for i in range(n):
            x1,y1=poly[i]; x2,y2=poly[(i+1)%n]; cross=x1*y2 - x2*y1; A+=cross; Cx+=(x1+x2)*cross; Cy+=(y1+y2)*cross
        A*=0.5
        if abs(A)<1e-9: xs,ys=zip(*poly); return (sum(xs)/n, sum(ys)/n)
        return (Cx/(6*A), Cy/(6*A))
    w=h=20; boxes=[]
    for k,poly in houses.items():
        x,y = centroid(poly); left = x - w/2; top = y - h/2; txt = labels[k]
        boxes.append(f'''
        <v:rect style="position:absolute;left:{left}pt;top:{top}pt;width:{w}pt;height:{h}pt;z-index:5" strokecolor="none">
          <v:textbox inset="0,0,0,0">
            <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:t>{txt}</w:t></w:r></w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>''')
    boxes_xml = "\\n".join(boxes)
    xml = f'''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r>
        <w:pict xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w10="urn:schemas-microsoft-com:office:word">
          <v:group style="position:relative;margin-left:0;margin-top:0;width:{S}pt;height:{S}pt" coordorigin="0,0" coordsize="{S},{S}">
            <v:rect style="position:absolute;left:0;top:0;width:{S}pt;height:{S}pt;z-index:1" strokecolor="black" strokeweight="1.5pt" fillcolor="#fff2cc"/>
            <v:line style="position:absolute;z-index:2" from="0,0" to="{S},{S}" strokecolor="black" strokeweight="1.5pt"/>
            <v:line style="position:absolute;z-index:2" from="{S},0" to="0,{S}" strokecolor="black" strokeweight="1.5pt"/>
            <v:line style="position:absolute;z-index:2" from="{S/2},0" to="{S},{S/2}" strokecolor="black" strokeweight="1.5pt"/>
            <v:line style="position:absolute;z-index:2" from="{S},{S/2}" to="{S/2},{S}" strokecolor="black" strokeweight="1.5pt"/>
            <v:line style="position:absolute;z-index:2" from="{S/2},{S}" to="0,{S/2}" strokecolor="black" strokeweight="1.5pt"/>
            <v:line style="position:absolute;z-index:2" from="0,{S/2}" to="{S/2},0" strokecolor="black" strokeweight="1.5pt"/>
            {boxes_xml}
          </v:group>
        </w:pict></w:r></w:p>'''
    return parse_xml(xml)

def add_table_borders(table, size=6):
    tbl = table._tbl; tblPr = tbl.tblPr; tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'),'single'); el.set(qn('w:sz'),str(size)); tblBorders.append(el)
    tblPr.append(tblBorders)

def set_table_font(table, pt=8.0):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(pt)

def center_header_row(table):
    for cell in table.rows[0].cells:
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if par.runs: par.runs[0].bold = True

def set_col_widths(table, widths_inch):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_inch):
            row.cells[i].width = Inches(w)

def sanitize_filename(name: str) -> str:
    cleaned = "".join(ch for ch in (name or "Horoscope") if ch.isalnum() or ch in "_- ")
    cleaned = cleaned.strip().replace(" ", "_")
    return cleaned or "Horoscope"

def _utc_to_local(dt_utc, tzname, tz_hours, used_manual):
    if used_manual: return dt_utc + datetime.timedelta(hours=tz_hours)
    try:
        tz = pytz.timezone(tzname); return tz.fromutc(dt_utc.replace(tzinfo=pytz.utc))
    except Exception:
        return dt_utc + datetime.timedelta(hours=tz_hours)

def main():
    st.title(APP_TITLE)
    col0, col1 = st.columns([1.2, 1])
    with col0:
        name = st.text_input("Name")
        dob = st.date_input("Date of Birth", min_value=datetime.date(1800,1,1), max_value=datetime.date(2100,12,31))
        tob = st.time_input("Time of Birth", step=datetime.timedelta(minutes=1))
        place = st.text_input("Place of Birth (City, State, Country)")
        tz_override = st.text_input("UTC offset override (optional, e.g., 5.5)", "")
    with col1:
        pass

    api_key = st.secrets.get("GEOAPIFY_API_KEY","")

    if st.button("Generate DOCX"):
        try:
            lat, lon, disp = geocode(place, api_key)
            dt_local = datetime.datetime.combine(dob, tob)
            used_manual = False
            if tz_override.strip():
                tz_hours = float(tz_override)
                dt_utc = dt_local - datetime.timedelta(hours=tz_hours)
                tzname = f"UTC{tz_hours:+.2f} (manual)"  # show ONLY this string later
                used_manual = True
            else:
                tzname, tz_hours, dt_utc = tz_from_latlon(lat, lon, dt_local)

            jd, ay, sidelons = sidereal_positions(dt_utc)
            lagna_sign, asc_sid = ascendant_sign(jd, lat, lon, ay)
            nav_lagna_sign = navamsa_sign_from_lon_sid(asc_sid)

            df_positions = positions_table_no_symbol(sidelons)

            md_segments_utc = build_mahadashas_days_utc(dt_utc, sidelons['Mo'])

            def age_years(birth_dt_local, end_utc):
                local_end = _utc_to_local(end_utc, tzname, tz_hours, used_manual)
                days = (local_end.date() - birth_dt_local.date()).days
                return int(days // YEAR_DAYS)

            df_md = pd.DataFrame([
                {"Planet": HN[s["planet"]],
                 "Start Date": _utc_to_local(s["start"], tzname, tz_hours, used_manual).strftime("%d-%m-%Y"),
                 "Age (years)": age_years(dt_local, s["end"])}
                for s in md_segments_utc
            ])

            now_utc = datetime.datetime.utcnow()
            rows_ap = next_ant_praty_in_days_utc(now_utc, md_segments_utc, days_window=2*365)
            df_ap = pd.DataFrame([
                {"Major Dasha": HN[r["major"]], "Antar Dasha": HN[r["antar"]],
                 "Pratyantar Dasha": HN[r["pratyantar"]],
                 "Date": _utc_to_local(r["end"], tzname, tz_hours, used_manual).strftime("%d-%m-%Y")}
                for r in rows_ap
            ])

            img_lagna = render_north_diamond(size_px=900, stroke=3)
            img_nav   = render_north_diamond(size_px=900, stroke=3)

            # DOCX
            doc = Document()
            sec = doc.sections[0]; sec.page_width = Mm(210); sec.page_height = Mm(297)
            margin = Mm(12); sec.left_margin = sec.right_margin = margin; sec.top_margin = Mm(10); sec.bottom_margin = Mm(10)

            style = doc.styles['Normal']; style.font.name = LATIN_FONT; style.font.size = Pt(BASE_FONT_PT)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), HINDI_FONT); style._element.rPr.rFonts.set(qn('w:cs'), HINDI_FONT)

            title = doc.add_paragraph(f"{name or '—'} — Horoscope"); title.runs[0].font.size = Pt(BASE_FONT_PT+3); title.runs[0].bold = True

            outer = doc.add_table(rows=1, cols=2); outer.autofit=False
            right_width_in = 3.3; outer.columns[0].width = Inches(3.3); outer.columns[1].width = Inches(right_width_in)
            add_table_borders(outer, size=6)

            left = outer.rows[0].cells[0]
            p = left.add_paragraph("Personal Details"); p.runs[0].bold=True
            left.add_paragraph(f"Name: {name}")
            left.add_paragraph(f"DOB: {dob}  |  TOB: {tob}")
            left.add_paragraph(f"Place: {disp}")
            # Fix timezone duplication:
            if used_manual:
                left.add_paragraph(f"Time Zone: {tzname}")   # e.g., UTC+5.50 (manual)
            else:
                left.add_paragraph(f"Time Zone: {tzname} (UTC{tz_hours:+.2f})")
            # Removed: Ayanamsha/Node/Year-basis line

            left.add_paragraph("Planetary Positions (sidereal, Swiss SWIEPH)").runs[0].bold=True
            t1 = left.add_table(rows=1, cols=len(df_positions.columns)); t1.autofit=False
            for i,c in enumerate(df_positions.columns): t1.rows[0].cells[i].text=c
            for _,row in df_positions.iterrows():
                r=t1.add_row().cells
                for i,c in enumerate(row): r[i].text=str(c)
            center_header_row(t1); set_table_font(t1, pt=BASE_FONT_PT); add_table_borders(t1, size=6)
            set_col_widths(t1, [0.75,0.5,0.9,0.85,0.85])

            left.add_paragraph("Vimshottari Mahadasha (start date + age in years)").runs[0].bold=True
            t2 = left.add_table(rows=1, cols=len(df_md.columns)); t2.autofit=False
            for i,c in enumerate(df_md.columns): t2.rows[0].cells[i].text=c
            for _,row in df_md.iterrows():
                r=t2.add_row().cells
                for i,c in enumerate(row): r[i].text=str(c)
            center_header_row(t2); set_table_font(t2, pt=BASE_FONT_PT); add_table_borders(t2, size=6)
            set_col_widths(t2, [0.9,0.95,0.9])  # tightened

            left.add_paragraph("Antar / Pratyantar (Next 2 years)").runs[0].bold=True
            t3 = left.add_table(rows=1, cols=len(df_ap.columns)); t3.autofit=False
            for i,c in enumerate(df_ap.columns): t3.rows[0].cells[i].text=c
            for _,row in df_ap.iterrows():
                r=t3.add_row().cells
                for i,c in enumerate(row): r[i].text=str(c)
            center_header_row(t3); set_table_font(t3, pt=BASE_FONT_PT); add_table_borders(t3, size=6)
            set_col_widths(t3, [0.85,0.9,1.0,0.75])  # tighter to avoid clipping

            right = outer.rows[0].cells[1]
            kt = right.add_table(rows=2, cols=1); kt.autofit=False; kt.columns[0].width = Inches(right_width_in)
            for row in kt.rows: row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY; row.height = Pt(340)

            cell1 = kt.rows[0].cells[0]; cell1.add_paragraph(); cap1 = cell1.add_paragraph("लग्न कुंडली")
            cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER; _apply_hindi_caption_style(cap1, size_pt=11, underline=True, bold=True)
            p1 = cell1.add_paragraph(); p1._p.addnext(kundali_w_p_with_centroid_labels(size_pt=220, lagna_sign=lagna_sign))

            cell2 = kt.rows[1].cells[0]; cell2.add_paragraph(); cap2 = cell2.add_paragraph("নवांश कुंडली".replace("ন","न"))
            cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER; _apply_hindi_caption_style(cap2, size_pt=11, underline=True, bold=True)
            p2 = cell2.add_paragraph(); p2._p.addnext(kundali_w_p_with_centroid_labels(size_pt=220, lagna_sign=nav_lagna_sign))

            out = BytesIO(); doc.save(out); out.seek(0)
            # File name now just "<Name>_Horoscope.docx"
            st.download_button("⬇️ Download DOCX", out.getvalue(), file_name=f"{sanitize_filename(name)}_Horoscope.docx")

            lc, rc = st.columns([1.2, 0.8])
            with lc:
                st.subheader("Planetary Positions")
                st.dataframe(df_positions.reset_index(drop=True), use_container_width=True, hide_index=True)
                st.subheader("Vimshottari Mahadasha (start date + age in years)")
                st.dataframe(df_md.reset_index(drop=True), use_container_width=True, hide_index=True)
                st.subheader("Antar / Pratyantar (Next 2 years)")
                st.dataframe(df_ap.reset_index(drop=True), use_container_width=True, hide_index=True)
            with rc:
                st.subheader("Lagna Kundali (Preview)")
                st.image(img_lagna, use_container_width=True)
                st.subheader("Navamsa Kundali (Preview)")
                st.image(img_nav, use_container_width=True)

        except Exception as e:
            st.error(str(e))

if __name__=='__main__':
    main()
